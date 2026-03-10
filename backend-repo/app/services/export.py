from fpdf import FPDF
from docx import Document
import os

def export_pdf(data: dict, output_path: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(40, 10, data.get('title', 'Meeting Report'))
    pdf.ln(10)
    
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, f"Summary:\n{data.get('summary', '')}")
    pdf.ln(5)
    pdf.multi_cell(0, 10, f"Transcript:\n{data.get('transcript', '')}")
    
    pdf.output(output_path)

def export_docx(data: dict, output_path: str):
    doc = Document()
    doc.add_heading(data.get('title', 'Meeting Report'), 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(data.get('summary', ''))
    
    doc.add_heading('Transcript', level=1)
    doc.add_paragraph(data.get('transcript', ''))
    
    doc.save(output_path)
